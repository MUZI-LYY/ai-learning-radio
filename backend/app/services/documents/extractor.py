"""文档正文提取：DOCX / 文本型 PDF / Markdown。

只支持 PRD 指定格式；不执行 HTML、脚本或远程资源。
"""

from __future__ import annotations

import io
import re

from app.core.errors import ErrorCode
from app.models.enums import MediaType

_WHITESPACE_RE = re.compile(r"[ \t\u00a0]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def _normalize(text: str) -> str:
    text = _WHITESPACE_RE.sub(" ", text)
    text = _BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return _normalize("\n".join(parts))


def _extract_pdf(data: bytes) -> tuple[str, str | None]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) == 0:
        return "", ErrorCode.EMPTY_DOCUMENT
    texts = [page.extract_text() or "" for page in reader.pages]
    text = _normalize("\n".join(texts))
    if len(text) < 20:
        return "", ErrorCode.SCANNED_PDF_UNSUPPORTED
    return text, None


def _extract_markdown(data: bytes) -> str:
    text = data.decode("utf-8")
    # 移除图片语法，保留标题层级与正文；不解析 HTML/脚本
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text)
    return _normalize(text)


def extract_text(media_type: MediaType, data: bytes) -> tuple[str, str | None]:
    """返回 (规范化正文, 错误码)；错误码为 None 表示成功。"""
    try:
        if media_type == MediaType.DOCX:
            text = _extract_docx(data)
            if not text:
                return "", ErrorCode.EMPTY_DOCUMENT
            return text, None
        if media_type == MediaType.PDF:
            return _extract_pdf(data)
        if media_type == MediaType.MARKDOWN:
            text = _extract_markdown(data)
            if not text:
                return "", ErrorCode.EMPTY_DOCUMENT
            return text, None
    except Exception:  # noqa: BLE001 - 解析异常统一视为不支持/不可解析
        return "", ErrorCode.UNSUPPORTED_FILE
    return "", ErrorCode.UNSUPPORTED_FILE
